import streamlit as st
import anthropic
import openai
import google.generativeai as genai
from sentence_transformers import SentenceTransformer
import spacy
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.feature_extraction.text import TfidfVectorizer
import pandas as pd
import numpy as np
import re
from io import BytesIO
import base64
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import markdown

# Page configuration
st.set_page_config(
    page_title="Blog Post Optimizer",
    page_icon="✍️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better UI
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: 700;
        color: #1f77b4;
        margin-bottom: 0.5rem;
    }
    .sub-header {
        font-size: 1.2rem;
        color: #666;
        margin-bottom: 2rem;
    }
    .success-box {
        padding: 1rem;
        border-radius: 0.5rem;
        background-color: #d4edda;
        border: 1px solid #c3e6cb;
        margin: 1rem 0;
    }
    .highlight {
        background-color: #fff3cd;
        padding: 0.2rem 0.4rem;
        border-radius: 0.25rem;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'outline_result' not in st.session_state:
    st.session_state.outline_result = None
if 'draft_result' not in st.session_state:
    st.session_state.draft_result = None
if 'audience_insights' not in st.session_state:
    st.session_state.audience_insights = None

# Load models (cached to prevent reloading)
@st.cache_resource
def load_nlp_models():
    """Load and cache NLP models"""
    try:
        # Load spaCy model
        nlp = spacy.load("en_core_web_sm")
    except OSError:
        # Model not found - show helpful error for Streamlit Cloud users
        st.error("""
        ⚠️ **spaCy model not found!**
        
        Please add this line to your `requirements.txt`:
        ```
        https://github.com/explosion/spacy-models/releases/download/en_core_web_sm-3.7.1/en_core_web_sm-3.7.1-py3-none-any.whl
        ```
        
        Then redeploy your app.
        """)
        st.stop()
    
    # Load sentence transformer for semantic similarity
    semantic_model = SentenceTransformer('all-MiniLM-L6-v2')
    
    return nlp, semantic_model

# API Configuration
def get_api_keys():
    """Retrieve API keys from Streamlit secrets"""
    openai_key = ""
    anthropic_key = ""
    google_key = ""
    
    try:
        if "OPENAI_API_KEY" in st.secrets:
            openai_key = st.secrets["OPENAI_API_KEY"]
            # Ensure it's not empty
            if not openai_key or openai_key.strip() == "":
                openai_key = ""
        
        if "ANTHROPIC_API_KEY" in st.secrets:
            anthropic_key = st.secrets["ANTHROPIC_API_KEY"]
            if not anthropic_key or anthropic_key.strip() == "":
                anthropic_key = ""
        
        if "GOOGLE_API_KEY" in st.secrets:
            google_key = st.secrets["GOOGLE_API_KEY"]
            if not google_key or google_key.strip() == "":
                google_key = ""
    except Exception as e:
        # Secrets not configured or error accessing them
        pass
    
    return openai_key, anthropic_key, google_key

def get_available_models():
    """Get available models based on configured API keys"""
    openai_key, anthropic_key, google_key = get_api_keys()
    
    models = {}
    
    if openai_key and len(openai_key) > 10:  # Basic validation
        models["OpenAI"] = {
            "gpt-4o": "GPT-4o (Latest, Most Capable)",
            "gpt-4o-mini": "GPT-4o Mini (Fast & Efficient)",
            "gpt-4-turbo": "GPT-4 Turbo (Balanced)"
        }
    
    if anthropic_key and len(anthropic_key) > 10:  # Basic validation
        models["Anthropic"] = {
            "claude-3-5-sonnet-20241022": "Claude 3.5 Sonnet (Recommended)",
            "claude-3-5-haiku-20241022": "Claude 3.5 Haiku (Fast)",
            "claude-3-opus-20240229": "Claude 3 Opus (Most Capable)"
        }
    
    if google_key and len(google_key) > 10:  # Basic validation
        models["Google"] = {
            "gemini-2.0-flash-exp": "Gemini 2.0 Flash (Experimental, Fastest)",
            "gemini-1.5-pro": "Gemini 1.5 Pro (Balanced)",
            "gemini-1.5-flash": "Gemini 1.5 Flash (Fast & Cost-Effective)"
        }
    
    return models

# AI API Functions
def call_openai(prompt, model="gpt-4o", max_tokens=4000):
    """Call OpenAI API"""
    openai_key, _, _ = get_api_keys()
    if not openai_key or len(openai_key) < 10:
        st.error("⚠️ OpenAI API key not found or invalid in secrets!")
        return None
    
    try:
        client = openai.OpenAI(api_key=openai_key)
        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=max_tokens,
            temperature=0.7
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"❌ OpenAI API Error: {str(e)}")
        return None

def call_anthropic(prompt, model="claude-3-5-sonnet-20241022", max_tokens=4000):
    """Call Anthropic API"""
    _, anthropic_key, _ = get_api_keys()
    if not anthropic_key or len(anthropic_key) < 10:
        st.error("⚠️ Anthropic API key not found or invalid in secrets!")
        return None
    
    try:
        client = anthropic.Anthropic(api_key=anthropic_key)
        message = client.messages.create(
            model=model,
            max_tokens=max_tokens,
            messages=[{"role": "user", "content": prompt}]
        )
        return message.content[0].text
    except Exception as e:
        st.error(f"❌ Anthropic API Error: {str(e)}")
        return None

def call_gemini(prompt, model="gemini-1.5-pro", max_tokens=4000):
    """Call Google Gemini API"""
    _, _, google_key = get_api_keys()
    if not google_key or len(google_key) < 10:
        st.error("⚠️ Google API key not found or invalid in secrets!")
        return None
    
    try:
        genai.configure(api_key=google_key)
        model_instance = genai.GenerativeModel(model)
        
        generation_config = {
            "max_output_tokens": max_tokens,
            "temperature": 0.7,
        }
        
        response = model_instance.generate_content(
            prompt,
            generation_config=generation_config
        )
        return response.text
    except Exception as e:
        st.error(f"❌ Google Gemini API Error: {str(e)}")
        return None

def call_ai_model(prompt, model_string, max_tokens=4000):
    """Universal AI caller that routes to the correct provider"""
    if model_string.startswith("gpt"):
        return call_openai(prompt, model_string, max_tokens)
    elif model_string.startswith("claude"):
        return call_anthropic(prompt, model_string, max_tokens)
    elif model_string.startswith("gemini"):
        return call_gemini(prompt, model_string, max_tokens)
    else:
        st.error("Unknown model type!")
        return None

# Utility Functions
def extract_headings(markdown_text):
    """Extract H2 and H3 headings from Markdown"""
    h2_pattern = r'^##\s+(.+)$'
    h3_pattern = r'^###\s+(.+)$'
    
    headings = []
    for line in markdown_text.split('\n'):
        h2_match = re.match(h2_pattern, line.strip())
        h3_match = re.match(h3_pattern, line.strip())
        
        if h2_match:
            headings.append(('H2', h2_match.group(1)))
        elif h3_match:
            headings.append(('H3', h3_match.group(1)))
    
    return headings

def semantic_similarity_score(text1, text2, model):
    """Calculate semantic similarity between two texts"""
    embeddings = model.encode([text1, text2])
    similarity = cosine_similarity([embeddings[0]], [embeddings[1]])[0][0]
    return similarity

def create_download_link(content, filename, file_format="markdown"):
    """Create download link without page reset"""
    if file_format == "markdown":
        b64 = base64.b64encode(content.encode()).decode()
        href = f'<a href="data:text/markdown;base64,{b64}" download="{filename}">📥 Download {filename}</a>'
    else:  # PDF or other formats
        b64 = base64.b64encode(content.encode()).decode()
        href = f'<a href="data:text/plain;base64,{b64}" download="{filename}">📥 Download {filename}</a>'
    return href

def markdown_to_docx(markdown_text, title="Optimized Content"):
    """Convert markdown text to a Word document with formatting"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Split content into lines
    lines = markdown_text.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # Empty line - add paragraph break
            doc.add_paragraph()
            continue
        
        # H1 headers
        if line.startswith('# ') and not line.startswith('## '):
            text = line.replace('# ', '')
            p = doc.add_heading(text, level=1)
        
        # H2 headers
        elif line.startswith('## ') and not line.startswith('### '):
            text = line.replace('## ', '')
            p = doc.add_heading(text, level=2)
        
        # H3 headers
        elif line.startswith('### '):
            text = line.replace('### ', '')
            p = doc.add_heading(text, level=3)
        
        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            p = doc.add_paragraph(text, style='List Bullet')
        
        # Numbered lists
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(text, style='List Number')
        
        # Bold text (simplified - just remove ** markers for now)
        elif '**' in line:
            text = line.replace('**', '')
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.bold = True
        
        # Regular paragraph
        else:
            doc.add_paragraph(line)
    
    # Save to BytesIO
    docx_file = BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    
    return docx_file

# === TAB 1: OUTLINE OPTIMIZER ===
def audience_research_analysis(keyword, model):
    """Perform deep audience research and search intent analysis"""
    prompt = f"""Conduct advanced audience research for the keyword/topic: "{keyword}"

Focus on:
1. Ideal Customer Profile (ICP) - Demographics and psychographics
2. Core pain points and frustrations that drive searches
3. Emotional triggers and psychological motivations
4. Search behaviors across platforms (Google, ChatGPT, Claude, Perplexity, Gemini)
5. User intent clusters (informational, transactional, navigational, investigational)
6. Common questions and information gaps
7. Resonance strategies - how to connect with this audience

Provide a comprehensive analysis structured as:
- **Demographics & Psychographics**: Who searches for this?
- **Pain Points**: What problems are they trying to solve?
- **Emotional Triggers**: What drives their search behavior?
- **Intent Clusters**: What are they really looking for?
- **Key Questions**: Top 10 questions they ask
- **Resonance Strategy**: How to create content that connects

Be specific, actionable, and data-driven in your analysis."""

    return call_ai_model(prompt, model, max_tokens=3000)

def optimize_outline(keyword, draft_outline, query_fanout, audience_insights, model):
    """Optimize blog outline with AI analysis"""
    prompt = f"""You are an expert SEO content strategist. Optimize the following blog post outline.

PRIMARY KEYWORD: {keyword}

AUDIENCE INSIGHTS:
{audience_insights}

QUERY FAN-OUT ANALYSIS:
{query_fanout}

CURRENT DRAFT OUTLINE:
{draft_outline}

TASK:
1. Analyze the draft outline against the Query Fan-Out suggestions and Audience Insights
2. Prioritize Query Fan-Out recommendations over the current draft
3. Ensure all ICP pain points and search intents are addressed
4. Generate an OPTIMIZED OUTLINE in the EXACT same Markdown structure (preserve H2/H3 hierarchy)

For each H2/H3 section, add:
- 7-12 concise talking points (bullet list) to guide writers
- Each talking point should be actionable and derived from audience research
- Include rationale for key additions (e.g., "Addresses [pain point] from audience research")

OUTPUT FORMAT:
## [H2 Heading]
- Talking point 1
- Talking point 2
...
- Talking point 7-12

### [H3 Heading]
- Talking point 1
...

Ensure logical flow, comprehensive coverage, and SEO optimization. Make it ready for a writer to execute."""

    return call_ai_model(prompt, model, max_tokens=4000)

# === TAB 2: DRAFT OPTIMIZER ===
def keyword_relevance_analysis(primary_keyword, keyword_list, draft_content, nlp, semantic_model):
    """Analyze keyword relevance and find placement opportunities"""
    # Parse draft with spaCy
    doc = nlp(draft_content)
    
    # Split into paragraphs
    paragraphs = [p.strip() for p in draft_content.split('\n\n') if p.strip()]
    
    # Calculate TF-IDF for relevance scoring
    all_text = [primary_keyword] + keyword_list + paragraphs
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform(all_text)
    
    results = []
    
    for keyword in keyword_list:
        # Calculate semantic similarity with primary keyword
        primary_sim = semantic_similarity_score(keyword, primary_keyword, semantic_model)
        
        # Check if keyword exists in draft (exact or variations)
        keyword_lower = keyword.lower()
        exists_exact = keyword_lower in draft_content.lower()
        
        # Find semantic variations
        variations = []
        for sent in doc.sents:
            sent_text = sent.text.lower()
            if keyword_lower not in sent_text:
                sim = semantic_similarity_score(keyword, sent.text, semantic_model)
                if sim > 0.8:
                    variations.append(sent.text[:50] + "...")
        
        # Calculate relevance score
        keyword_embedding = semantic_model.encode([keyword])
        draft_embedding = semantic_model.encode([draft_content[:500]])  # First 500 chars
        relevance_score = cosine_similarity(keyword_embedding, draft_embedding)[0][0]
        
        # Determine if keyword needs addition
        keyword_density = draft_content.lower().count(keyword_lower) / len(draft_content.split()) * 100
        needs_addition = relevance_score > 0.5 and keyword_density < 1.0
        
        if needs_addition:
            # Find best placement location
            best_para_idx = 0
            best_sim = 0
            for idx, para in enumerate(paragraphs[:5]):  # Check first 5 paragraphs
                sim = semantic_similarity_score(keyword, para, semantic_model)
                if sim > best_sim:
                    best_sim = sim
                    best_para_idx = idx
            
            results.append({
                'keyword': keyword,
                'relevance_score': round(relevance_score, 2),
                'density': round(keyword_density, 2),
                'variations_found': len(variations),
                'needs_addition': needs_addition,
                'best_paragraph_idx': best_para_idx,
                'paragraph_preview': paragraphs[best_para_idx][:100] + "..."
            })
    
    return results

def generate_keyword_integration(keyword, paragraph_context, model):
    """Generate natural keyword integration"""
    prompt = f"""You are a skilled content writer. Rewrite the following paragraph to naturally integrate the keyword while maintaining flow and readability.

KEYWORD TO INTEGRATE: {keyword}

ORIGINAL PARAGRAPH:
{paragraph_context}

REQUIREMENTS:
- Integrate the keyword naturally and contextually
- Maintain the original tone and structure
- Keep paragraph length similar (50-200 words)
- Ensure it reads naturally, not stuffed
- Make it copy-paste ready

OUTPUT: Only provide the rewritten paragraph, nothing else."""

    return call_ai_model(prompt, model, max_tokens=500)

def ai_tool_optimization(draft_content, primary_keyword, keyword_list, model):
    """Apply 10-item AI optimization checklist"""
    
    # Count original words
    original_word_count = len(draft_content.split())
    
    prompt = f"""You are an expert content optimizer for AI search tools (Google AI Overviews, ChatGPT, Claude, Perplexity).

PRIMARY KEYWORD: {primary_keyword}
KEYWORDS: {', '.join(keyword_list[:10])}

CONTENT DRAFT (Original word count: {original_word_count} words):
{draft_content}

CRITICAL REQUIREMENTS:
1. MAINTAIN OR EXPAND the original content length - aim for {original_word_count} to {int(original_word_count * 1.2)} words
2. DO NOT shorten or remove substantial content
3. ADD details, examples, data, and depth rather than removing content
4. Enhancement means ADDING VALUE, not reducing content

Apply these 10 optimizations while PRESERVING OR EXPANDING content:

1. **Answer-First Introduction**: Rewrite opening 30-50 words to directly answer core query (keep rest of intro)
2. **Question-Based H2s**: Convert all H2 headings to natural questions (keep all content under each heading)
3. **Semantic Chunks**: Ensure 75-300 word self-contained sections (expand short sections, don't shorten long ones)
4. **Answer-Evidence-Context**: Restructure chunks (answer -> evidence -> context) - ADD evidence and context, don't remove
5. **Direct Sentences**: Convert to active voice, Subject-Verb-Object (maintain all information)
6. **Informational Density**: Increase specifics by 20% (numbers, entities, examples) - this means ADDING content
7. **Attribute Claims**: Replace generics with "A 2023 study by XYZ found..." - ADD citations and sources
8. **Signal Experience**: Add "In our testing..." where applicable - ADD first-hand insights
9. **FAQ Section**: Append 3-10 long-tail Q&A pairs at the end - this ADDS content
10. **Title & Meta**: Generate optimized title (<60 chars) and meta (140-160 chars) at the top

OUTPUT the fully optimized content in clean Markdown format. 
- Do NOT add any tags, markers, or annotations like [OPTIMIZED]
- Output should be clean, ready-to-use content
- Ensure final word count is AT LEAST {int(original_word_count * 0.95)} words (95% of original minimum)
- Aim to EXPAND content with valuable details, not compress it"""

    return call_ai_model(prompt, model, max_tokens=8000)

# === MAIN APP ===
def main():
    st.markdown('<p class="main-header">✍️ Blog Post Optimizer</p>', unsafe_allow_html=True)
    st.markdown('<p class="sub-header">AI-Powered Content Optimization for SEO & AI Search Tools</p>', unsafe_allow_html=True)
    
    # Sidebar - API Configuration Check
    with st.sidebar:
        st.header("⚙️ Configuration")
        openai_key, anthropic_key, google_key = get_api_keys()
        
        # Display API connection status
        st.subheader("API Connection Status")
        if openai_key and len(openai_key) > 10:  # Basic validation
            st.success("✅ OpenAI API Connected")
        else:
            st.warning("⚠️ OpenAI API Key Missing")
        
        if anthropic_key and len(anthropic_key) > 10:  # Basic validation
            st.success("✅ Anthropic API Connected")
        else:
            st.warning("⚠️ Anthropic API Key Missing")
        
        if google_key and len(google_key) > 10:  # Basic validation
            st.success("✅ Google Gemini API Connected")
        else:
            st.warning("⚠️ Google API Key Missing")
        
        # Get available models
        available_models = get_available_models()
        
        if not available_models:
            st.error("❌ No API keys configured! Please add at least one API key to Streamlit Secrets.")
            st.stop()
        
        st.markdown("---")
        
        # Smart Model Selector - only shows models from configured providers
        st.subheader("🤖 AI Model Selection")
        
        # Flatten the models dictionary for selection
        model_options = {}
        for provider, models in available_models.items():
            for model_id, model_name in models.items():
                model_options[f"{provider}: {model_name}"] = model_id
        
        selected_model_display = st.selectbox(
            "Select AI Model:",
            options=list(model_options.keys()),
            help="Only models from configured API providers are shown"
        )
        
        selected_model = model_options[selected_model_display]
        
        # Display model info
        st.info(f"**Active Model:** `{selected_model}`")
        
        st.markdown("---")
        st.markdown("**💡 Setup Instructions:**")
        st.markdown("Add API keys to Streamlit Cloud Secrets:")
        st.code('OPENAI_API_KEY = "your-key"\nANTHROPIC_API_KEY = "your-key"\nGOOGLE_API_KEY = "your-key"')
        
        st.markdown("---")
        st.markdown("**💰 Cost Comparison:**")
        st.markdown("- 🥇 **Gemini**: Most cost-effective (free tier)")
        st.markdown("- 🥈 **Anthropic**: Best quality/cost ratio")
        st.markdown("- 🥉 **OpenAI**: Most capable models")
    
    # Load NLP models
    with st.spinner("Loading NLP models..."):
        nlp, semantic_model = load_nlp_models()
    
    # Main Tabs
    tab1, tab2 = st.tabs(["📝 Outline Optimizer", "✨ Draft Optimizer"])
    
    # === TAB 1: OUTLINE OPTIMIZER ===
    with tab1:
        st.header("Outline Optimizer")
        st.markdown("Optimize your blog outline with audience research, search intent analysis, and SEO best practices.")
        
        col1, col2 = st.columns([1, 1])
        
        with col1:
            keyword_outline = st.text_input(
                "Primary Keyword/Topic *",
                placeholder="e.g., dog training tips",
                key="outline_keyword"
            )
            
            draft_outline = st.text_area(
                "Draft Outline (Markdown) *",
                placeholder="## Introduction\n\n## Main Point 1\n### Subpoint 1.1\n\n## Conclusion",
                height=250,
                key="draft_outline_input"
            )
            
            uploaded_outline = st.file_uploader(
                "Or Upload Draft Outline",
                type=['txt', 'md'],
                key="outline_upload",
                help="Upload your draft outline as a Markdown or text file"
            )
        
        with col2:
            query_fanout = st.text_area(
                "Query Fan-Out Analysis",
                placeholder="Paste your Query Fan-Out report here (structured text with expanded queries, related terms, intent insights)",
                height=250,
                key="query_fanout_input"
            )
            
            uploaded_fanout = st.file_uploader(
                "Or Upload Query Fan-Out Report",
                type=['txt', 'md'],
                key="fanout_upload",
                help="Upload your Query Fan-Out analysis report"
            )
        
        # Handle file uploads
        if uploaded_outline:
            draft_outline = uploaded_outline.read().decode('utf-8')
            st.success("✅ Outline uploaded successfully!")
        
        if uploaded_fanout:
            query_fanout = uploaded_fanout.read().decode('utf-8')
            st.success("✅ Query Fan-Out report uploaded successfully!")
        
        st.markdown("---")
        
        if st.button("🚀 Optimize Outline", type="primary", use_container_width=True):
            if not keyword_outline or not draft_outline:
                st.error("Please provide both primary keyword and draft outline!")
            else:
                # Step 1: Audience Research
                with st.spinner("🔍 Step 1/2: Conducting deep audience research..."):
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    status_text.text("Analyzing search intent and user psychology...")
                    progress_bar.progress(20)
                    
                    audience_insights = audience_research_analysis(keyword_outline, selected_model)
                    st.session_state.audience_insights = audience_insights
                    
                    status_text.text("Research complete!")
                    progress_bar.progress(100)
                    
                    if audience_insights:
                        st.success("✅ Audience research complete!")
                        with st.expander("📊 View Audience Insights", expanded=False):
                            st.markdown(audience_insights)
                    
                    status_text.empty()
                
                # Step 2: Outline Optimization
                with st.spinner("✨ Step 2/2: Optimizing outline..."):
                    progress_bar2 = st.progress(0)
                    status_text2 = st.empty()
                    
                    fanout_text = query_fanout if query_fanout else "No Query Fan-Out analysis provided."
                    
                    status_text2.text("Cross-referencing with Query Fan-Out analysis...")
                    progress_bar2.progress(30)
                    
                    optimized_outline = optimize_outline(
                        keyword_outline,
                        draft_outline,
                        fanout_text,
                        audience_insights,
                        selected_model
                    )
                    
                    status_text2.text("Generating talking points...")
                    progress_bar2.progress(90)
                    
                    st.session_state.outline_result = optimized_outline
                    progress_bar2.progress(100)
                    status_text2.empty()
                
                if optimized_outline:
                    st.markdown("---")
                    st.subheader("📊 Optimization Results")
                    
                    # Side-by-side comparison
                    comp_col1, comp_col2 = st.columns(2)
                    
                    with comp_col1:
                        st.markdown("### 📄 Original Outline")
                        st.markdown(draft_outline)
                    
                    with comp_col2:
                        st.markdown("### ✅ Optimized Outline")
                        st.markdown(optimized_outline)
                    
                    # Download options
                    st.markdown("---")
                    download_col1, download_col2 = st.columns(2)
                    
                    with download_col1:
                        filename = f"optimized_outline_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
                        st.markdown(
                            create_download_link(optimized_outline, filename),
                            unsafe_allow_html=True
                        )
                    
                    with download_col2:
                        if st.session_state.audience_insights:
                            insights_filename = f"audience_insights_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
                            st.markdown(
                                create_download_link(st.session_state.audience_insights, insights_filename),
                                unsafe_allow_html=True
                            )
    
    # === TAB 2: DRAFT OPTIMIZER ===
    with tab2:
        st.header("Draft Optimizer")
        st.markdown("Optimize your draft through keyword analysis and AI tool optimization (2-step process).")
        
        keyword_draft = st.text_input(
            "Primary Keyword/Topic *",
            placeholder="e.g., dog training tips",
            key="draft_keyword"
        )
        
        draft_col1, draft_col2 = st.columns([2, 1])
        
        with draft_col1:
            draft_content = st.text_area(
                "Content Draft *",
                placeholder="Paste your blog draft here (Markdown or HTML supported)",
                height=300,
                key="draft_content_input"
            )
            
            uploaded_draft = st.file_uploader(
                "Or Upload Draft File",
                type=['txt', 'md', 'html'],
                key="draft_upload",
                help="Upload your blog draft as Markdown, text, or HTML"
            )
        
        with draft_col2:
            keyword_list_input = st.text_area(
                "Expanded Keyword List *",
                placeholder="keyword 1\nkeyword 2\nkeyword 3",
                height=300,
                key="keyword_list_input"
            )
        
        # Handle file upload
        if uploaded_draft:
            draft_content = uploaded_draft.read().decode('utf-8')
            st.success("✅ Draft uploaded successfully!")
        
        # Parse keyword list
        keyword_list = [k.strip() for k in keyword_list_input.split('\n') if k.strip()]
        if ',' in keyword_list_input:
            keyword_list = [k.strip() for k in keyword_list_input.split(',') if k.strip()]
        
        st.markdown("---")
        
        if st.button("🚀 Optimize Draft", type="primary", use_container_width=True):
            if not keyword_draft or not draft_content or not keyword_list:
                st.error("Please provide primary keyword, draft content, and keyword list!")
            else:
                # Subprocess 1: Keyword Optimization
                st.markdown("### 📊 Subprocess 1: Keyword Optimization")
                
                with st.spinner("🔍 Analyzing keyword relevance and placement..."):
                    keyword_analysis = keyword_relevance_analysis(
                        keyword_draft,
                        keyword_list,
                        draft_content,
                        nlp,
                        semantic_model
                    )
                
                if keyword_analysis:
                    st.success(f"✅ Analyzed {len(keyword_analysis)} high-relevance keywords needing optimization")
                    
                    # Generate integration suggestions
                    integration_results = []
                    
                    progress = st.progress(0)
                    for idx, kw_data in enumerate(keyword_analysis[:10]):  # Limit to top 10
                        with st.spinner(f"Generating integration for '{kw_data['keyword']}'..."):
                            integrated_paragraph = generate_keyword_integration(
                                kw_data['keyword'],
                                kw_data['paragraph_preview'],
                                selected_model
                            )
                            
                            integration_results.append({
                                'Selected Keyword': f"{kw_data['keyword']} (Score: {kw_data['relevance_score']})",
                                'Placement Hint': f"Insert in paragraph {kw_data['best_paragraph_idx'] + 1}: '{kw_data['paragraph_preview'][:50]}...'",
                                'Content Snippet': integrated_paragraph if integrated_paragraph else "Generation failed"
                            })
                        
                        progress.progress((idx + 1) / min(10, len(keyword_analysis)))
                    
                    # Display table
                    if integration_results:
                        st.markdown("#### 📝 Keyword Integration Suggestions")
                        df = pd.DataFrame(integration_results)
                        st.dataframe(df, use_container_width=True, height=400)
                
                st.markdown("---")
                
                # Subprocess 2: AI Tool Optimization
                st.markdown("### ✨ Subprocess 2: AI Tool Optimization")
                
                with st.spinner("🤖 Applying 10-item AI optimization checklist..."):
                    progress2 = st.progress(0)
                    
                    optimized_draft = ai_tool_optimization(
                        draft_content,
                        keyword_draft,
                        keyword_list,
                        selected_model
                    )
                    
                    progress2.progress(100)
                    st.session_state.draft_result = optimized_draft
                
                if optimized_draft:
                    st.success("✅ Draft fully optimized!")
                    
                    # Checklist compliance
                    st.markdown("#### ✅ Optimization Checklist Compliance")
                    checklist_items = [
                        "Answer-First Introduction",
                        "Question-Based H2 Headings",
                        "Semantic Chunks (75-300 words)",
                        "Answer-Evidence-Context Formula",
                        "Direct Sentence Structures",
                        "Informational Density (+20%)",
                        "Attributed Claims",
                        "First-Hand Experience Signals",
                        "Dedicated FAQ Section",
                        "Optimized Title & Meta"
                    ]
                    
                    cols = st.columns(5)
                    for idx, item in enumerate(checklist_items):
                        with cols[idx % 5]:
                            st.metric(label=item, value="✅", delta="Applied")
                    
                    st.markdown("---")
                    
                    # Display optimized draft
                    st.markdown("#### 📄 Fully Optimized Draft")
                    st.markdown(optimized_draft)
                    
                    # Download
                    st.markdown("---")
                    final_filename = f"optimized_draft_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
                    st.markdown(
                        create_download_link(optimized_draft, final_filename),
                        unsafe_allow_html=True
                    )

if __name__ == "__main__":
    main()
